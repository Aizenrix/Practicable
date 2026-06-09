#!/usr/bin/env python3
"""Генерация материалов УП.03 Этап 3 для Calipso.Coffee."""

import shutil
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

BASE = Path(__file__).resolve().parent.parent
OUT = BASE / "УП03_Этап3_Васильев_Данил_РПО-23-1"
REPO = "https://github.com/Aizenrix/Practicable.git"
PUBLIC_URL = "http://localhost:4000"
TODAY = date.today().strftime("%d.%m.%Y")

STUDENT_FIO = "Васильев Данил Вячеславович"
GROUP = "РПО 23/1"
SPECIALTY = "09.02.07 Информационные системы и программирование"
MODULE = "ПМ.03. Сопровождение и обслуживание программного обеспечения компьютерных систем"
STAGE = "Этап 3. Настройка компонентов и развертывание проекта"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("КОРОТКИЙ ОТЧЁТ ПО ЭТАПУ 3\nУП.03. Учебная практика\n")
    r.bold = True
    r.font.size = Pt(14)

    for text in [
        STAGE,
        "",
        "Проект: Calipso.Coffee",
        f"Студент: {STUDENT_FIO}, группа {GROUP}",
        f"Специальность: {SPECIALTY}",
        f"Дата: {TODAY}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).font.size = Pt(12)

    doc.add_page_break()

    doc.add_heading("1. Ссылка на репозиторий", level=1)
    doc.add_paragraph(f"Репозиторий: {REPO}")
    doc.add_paragraph("Ветка: main")

    doc.add_heading("2. Выбранный вариант развертывания", level=1)
    doc.add_paragraph(
        "Вариант C — локальный demo/production-стенд через Docker (docker-compose.prod.yml)."
    )
    doc.add_paragraph(
        "Причина выбора: проект — backend + SQLite + веб-интерфейс; учебный VPS недоступен, "
        "но Docker позволяет развернуть сервис вне IDE с restart policy, volume для БД, "
        "health-check и перезапуском — это соответствует требованиям этапа 3."
    )
    doc.add_paragraph(
        "Дополнительно подготовлены конфиги для VPS: nginx/project.conf, systemd/calipso-coffee.service."
    )

    doc.add_heading("3. Где открыть результат", level=1)
    doc.add_paragraph(f"Адрес demo-стенда: {PUBLIC_URL}")
    doc.add_paragraph("Health-check: http://localhost:4000/health")
    doc.add_paragraph("Release-архив: release/project_release.zip")

    doc.add_heading("4. Команды развертывания", level=1)
    for c in [
        "cp .env.production.example .env.production",
        "docker compose -f docker-compose.prod.yml up --build -d",
        "docker compose -f docker-compose.prod.yml ps",
        "curl http://localhost:4000/health",
        "docker compose -f docker-compose.prod.yml logs --tail=80",
        "make deploy-restart  # перезапуск",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("5. Переменные окружения", level=1)
    add_table(
        doc,
        ["Файл", "Назначение"],
        [
            [".env.production.example", "Production-конфигурация без секретов"],
            [".env.demo.example", "Демо-стенд, DEMO_MODE=true"],
            [".env.production", "Рабочий файл (не в Git)"],
        ],
    )
    add_table(
        doc,
        ["Переменная", "Пример", "Описание"],
        [
            ["APP_ENV", "production", "Режим приложения"],
            ["APP_DEBUG", "false", "Отключение debug"],
            ["PORT", "4000", "Порт HTTP"],
            ["JWT_SECRET", "change_me...", "Секрет JWT"],
            ["DATABASE_URL", "file:/data/dev.db", "SQLite в Docker volume"],
        ],
    )

    doc.add_heading("6. Проверка доступности", level=1)
    for item in [
        "Docker-контейнер calipso-coffee-prod в статусе running",
        "http://localhost:4000 открывается в браузере",
        "/health возвращает status: ok",
        "Вход admin@calipso.coffee / admin123 — успешен",
        "Разделы Сводка и Заказы загружаются",
        "Логи без критических ошибок",
        "Перезапуск (down + up -d) проходит успешно",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Проблемы и исправления", level=1)
    add_table(
        doc,
        ["№", "Проблема", "Решение"],
        [
            ["1", "Нет production-конфигурации", "Созданы .env.production.example, docker-compose.prod.yml"],
            ["2", "Проект доступен только через npm run dev", "Добавлены deploy/restart/check_deploy скрипты"],
            ["3", "EADDRINUSE порт 4000", "docker compose down перед повторным deploy"],
            ["4", "Нет release-пакета", "scripts/build_release.sh → release/project_release.zip"],
        ],
    )

    doc.add_heading("8. Скриншоты", level=1)
    add_table(
        doc,
        ["Файл", "Содержание"],
        [
            ["01_env_production_or_demo.png", ".env.production.example / .env.demo.example"],
            ["02_deploy_files.png", "DEPLOYMENT.md, docker-compose.prod.yml, scripts"],
            ["03_deploy_command_success.png", "make deploy или docker compose prod up -d"],
            ["04_service_or_build_started.png", "docker ps — контейнер running"],
            ["05_app_available.png", "Браузер http://localhost:4000"],
            ["06_logs_without_critical_errors.png", "docker compose logs"],
            ["07_main_scenario_works.png", "Вход и работа разделов"],
            ["08_restart_success.png", "make deploy-restart"],
            ["09_release_or_public_url.png", "public_url или release/project_release.zip"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph(f"Студент: {STUDENT_FIO}, группа {GROUP}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Дата: {TODAY}").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    (OUT / "docs").mkdir(parents=True, exist_ok=True)
    doc.save(OUT / "docs" / "01_Короткий_отчет_по_этапу_3.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    for sub in ["docs", "screenshots", "scripts", "release", "release/screenshots"]:
        (OUT / sub).mkdir(parents=True, exist_ok=True)

    (OUT / "repo_link.txt").write_text(REPO + "\n", encoding="utf-8")
    (OUT / "public_url.txt").write_text(
        f"{PUBLIC_URL}\n\nHealth: {PUBLIC_URL}/health\n"
        f"Release: release/project_release.zip\n",
        encoding="utf-8",
    )

    for name in ["DEPLOYMENT.md", "DEMO_GUIDE.md", "RELEASE_NOTES.md"]:
        shutil.copy(BASE / name, OUT / "docs" / name)

    for bat in ["deploy.bat", "restart.bat", "check_deploy.bat", "build_release.bat"]:
        shutil.copy(BASE / "scripts" / bat, OUT / "scripts" / bat)

    for rel in ["demo_readme.txt", "demo_accounts.txt"]:
        src = BASE / "release" / rel
        if src.exists():
            shutil.copy(src, OUT / "release" / rel)

    zip_src = BASE / "release" / "project_release.zip"
    if zip_src.exists():
        shutil.copy(zip_src, OUT / "release" / "project_release.zip")

    (OUT / "screenshots" / "README_скриншоты.txt").write_text(
        """Обязательные скриншоты этапа 3:

01_env_production_or_demo.png
02_deploy_files.png
03_deploy_command_success.png
04_service_or_build_started.png
05_app_available.png
06_logs_without_critical_errors.png
07_main_scenario_works.png
08_restart_success.png
09_release_or_public_url.png
""",
        encoding="utf-8",
    )

    (OUT / "README.md").write_text(
        f"""# УП.03 Этап 3 — Calipso.Coffee

**Студент:** {STUDENT_FIO}  
**Группа:** {GROUP}

## Адрес demo-стенда

{PUBLIC_URL}

## Команды

```bash
make deploy
make check-deploy
make deploy-restart
make build-release
```

## Состав

- docs/ — отчёт, DEPLOYMENT, DEMO_GUIDE, RELEASE_NOTES
- screenshots/ — 9 скриншотов
- release/ — demo_readme, demo_accounts, project_release.zip
- scripts/ — deploy, restart, check_deploy, build_release
""",
        encoding="utf-8",
    )

    build_report()
    print(f"Готово: {OUT}")


if __name__ == "__main__":
    main()
