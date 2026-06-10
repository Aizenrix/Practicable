#!/usr/bin/env python3
"""Генерация материалов УП.03 Этап 5 для Calipso.Coffee."""

import shutil
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

BASE = Path(__file__).resolve().parent.parent
OUT = BASE / "УП03_Этап5_Васильев_Данил_РПО-23-1"
REPO = "https://github.com/Aizenrix/Practicable.git"
PUBLIC_URL = "http://localhost:4000"
TODAY = date.today().strftime("%d.%m.%Y")

STUDENT_FIO = "Васильев Данил Вячеславович"
GROUP = "РПО 23/1"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build_security_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SECURITY REPORT — Этап 5\nУП.03. Безопасность и эксплуатационные риски\n")
    r.bold = True
    r.font.size = Pt(14)

    for text in [
        "Проект: Calipso.Coffee",
        f"Студент: {STUDENT_FIO}, группа {GROUP}",
        f"Дата: {TODAY}",
        f"Адрес: {PUBLIC_URL}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)

    doc.add_page_break()

    doc.add_heading("1. Что проверялось", level=1)
    for item in [
        "Секреты: .gitignore, .env.example, git grep secret scan",
        "Зависимости: npm audit, npm outdated",
        "Роли: ADMIN vs WAITER на /api/users",
        "Чужие данные: GET/PATCH /api/orders/:id — 403 для чужого заказа",
        "CORS: CORS_ORIGIN в production, APP_DEBUG=false",
        "Backup/restore SQLite: scripts/backup.sh, restore.sh",
        "Порты: docker compose ps, lsof/netstat",
        "Логи: docker logs без критических ошибок",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Инструменты", level=1)
    add_table(
        doc,
        ["Инструмент", "Назначение"],
        [
            ["git grep / security-check", "Поиск секретов"],
            ["npm audit", "Уязвимости зависимостей"],
            ["curl + JWT", "Проверка ролей и 403"],
            ["backup.sh / restore.sh", "Резервное копирование SQLite"],
            ["docker compose ps", "Порты и контейнеры"],
            ["node:test security/", "Автотесты доступа"],
        ],
    )

    doc.add_heading("3. Что найдено", level=1)
    add_table(
        doc,
        ["ID", "Проблема", "Критичность"],
        [
            ["SEC-01", "JWT_SECRET=change_me в примерах", "Средняя (demo)"],
            ["SEC-02", "npm audit: 2 moderate (qs)", "Средняя"],
            ["SEC-03", "Нет проверки владельца заказа (до этапа 5)", "Высокая → исправлено"],
            ["SEC-04", "CORS * по умолчанию в dev", "Низкая (prod настроен)"],
        ],
    )

    doc.add_heading("4. Что исправлено", level=1)
    for item in [
        "Добавлены backups/, *.sql в .gitignore",
        "canAccessOrder() — 403 на чужие заказы",
        "CORS_ORIGIN для production",
        "Скрипты backup/restore/security-check",
        "SECURITY.md, SECURITY_CHECKLIST.md",
        "Демо-пользователь waiter@calipso.coffee для проверки ролей",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Оставшиеся риски", level=1)
    doc.add_paragraph(
        "Нет HTTPS на localhost demo; npm audit moderate не закрыт полностью; "
        "нет автоматического cron-backup; нет rate limit на login. "
        "Подробнее: RISK_REGISTER.md, RECOMMENDATIONS.md."
    )

    doc.add_heading("6. Рекомендации", level=1)
    for item in [
        "CI: npm audit + secret scan на каждый push",
        "HTTPS через Nginx на VPS",
        "Ежедневный backup volume",
        "npm audit fix и мониторинг health",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Вывод", level=1)
    doc.add_paragraph(
        "Базовая проверка безопасности Calipso.Coffee выполнена: секреты не в Git, "
        "роли и доступ к чужим данным проверены, backup/restore SQLite работает, "
        "порты и логи задокументированы. Критическая проблема доступа к чужим заказам "
        "исправлена. Проект готов к сдаче этапа 5 с учётом оставшихся рисков S-04 и S-03."
    )

    doc.add_paragraph()
    p = doc.add_paragraph(f"Студент: {STUDENT_FIO}, группа {GROUP}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    (OUT / "docs").mkdir(parents=True, exist_ok=True)
    doc.save(OUT / "docs" / "SECURITY_REPORT.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    for sub in ["docs", "reports", "screenshots", "scripts", "project_files_in_repository"]:
        (OUT / sub).mkdir(parents=True, exist_ok=True)

    (OUT / "repo_link.txt").write_text(REPO + "\n", encoding="utf-8")
    (OUT / "public_url.txt").write_text(PUBLIC_URL + "\n", encoding="utf-8")

    for name in [
        "SECURITY_CHECKLIST.md",
        "RISK_REGISTER.md",
        "BACKUP_RESTORE_REPORT.md",
        "RECOMMENDATIONS.md",
        "SECURITY.md",
    ]:
        src = BASE / name
        if src.exists():
            shutil.copy(src, OUT / "docs" / name)

    report_files = [
        "security_secret_scan.txt",
        "security_deps_audit.txt",
        "security_backup.txt",
        "security_restore.txt",
        "security_ports.txt",
        "security_logs.txt",
        "security_access_check.txt",
    ]
    for rf in report_files:
        src = BASE / "reports" / rf
        if src.exists():
            shutil.copy(src, OUT / "reports" / rf)

    for bat in [
        "security-check.bat",
        "deps-check.bat",
        "backup.bat",
        "restore.bat",
        "ports-check.bat",
        "logs-security.bat",
        "role-access-check.bat",
    ]:
        src = BASE / "scripts" / bat
        if src.exists():
            shutil.copy(src, OUT / "scripts" / bat)

    proj_refs = [
        ".gitignore",
        ".env.example",
        ".env.production.example",
        "SECURITY.md",
        "docker-compose.yml",
        "docker-compose.prod.yml",
    ]
    ref_dir = OUT / "project_files_in_repository"
    (ref_dir / "README.txt").write_text(
        "Ключевые файлы этапа 5 в корне репозитория:\n"
        + "\n".join(f"  - {f}" for f in proj_refs)
        + "\n  - scripts/security-check.sh, backup.sh, restore.sh, ...\n",
        encoding="utf-8",
    )
    for f in proj_refs:
        src = BASE / f
        if src.exists():
            shutil.copy(src, ref_dir / Path(f).name)

    (OUT / "screenshots" / "README_скриншоты.txt").write_text(
        """Скриншоты этапа 5 (12 шт.):

01_gitignore_no_env.png          — .gitignore: .env, backups, logs, dumps
02_env_example_without_secrets.png — JWT_SECRET=change_me, без реальных ключей
03_secret_scan_result.png        — make security-check / git grep
04_dependency_check_result.png   — make deps-check / npm audit
05_roles_access_check.png        — WAITER 403 на /api/users, ADMIN 200
06_foreign_data_access_denied.png — GET /api/orders/:id чужого = 403
07_cors_or_security_config.png   — .env.production.example CORS_ORIGIN, app.js
08_backup_created.png            — backups/backup_*.db или make backup
09_restore_success.png           — make restore + health 200
10_open_ports_check.png          — docker ps / make ports-check
11_logs_no_critical_errors.png   — make logs-security
12_security_commit.png           — git log этапа 5
""",
        encoding="utf-8",
    )

    (OUT / "README.md").write_text(
        f"""# УП.03 Этап 5 — Calipso.Coffee

**Студент:** {STUDENT_FIO}  
**Группа:** {GROUP}

## Команды

```bash
make security-check
make deps-check
make backup
make restore
make ports-check
make logs-security
npm test
```

## Состав

- docs/ — SECURITY_REPORT.docx, CHECKLIST, RISK_REGISTER, BACKUP_RESTORE, RECOMMENDATIONS
- reports/ — secret scan, deps audit, backup, restore, ports, logs
- screenshots/ — 12 скриншотов
- scripts/ — security-check, deps-check, backup, restore, ports-check, logs-security
- project_files_in_repository/ — ссылки на ключевые файлы репозитория
""",
        encoding="utf-8",
    )

    build_security_report()
    print(f"Готово: {OUT}")


if __name__ == "__main__":
    main()
