#!/usr/bin/env python3
"""Генерация материалов УП.03 Этап 2 для Calipso.Coffee."""

from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

BASE = Path(__file__).resolve().parent.parent
OUT = BASE / "УП03_Этап2_Васильев_Данил_РПО-23-1"
REPO = "https://github.com/Aizenrix/Practicable.git"
TODAY = date.today().strftime("%d.%m.%Y")

STUDENT_FIO = "Васильев Данил Вячеславович"
GROUP = "РПО 23/1"
SPECIALTY = "09.02.07 Информационные системы и программирование"
MODULE = "ПМ.03. Сопровождение и обслуживание программного обеспечения компьютерных систем"
STAGE = "Этап 2. Проведение инсталляции ПО компьютерных систем"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    # Титул
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("КОРОТКИЙ ОТЧЁТ ПО ЭТАПУ 2\nУП.03. Учебная практика\n")
    r.bold = True
    r.font.size = Pt(14)

    for text in [
        STAGE,
        "",
        "Проект: Calipso.Coffee",
        "Система учёта заказов кафе",
        "",
        f"Студент: {STUDENT_FIO}",
        f"Группа: {GROUP}",
        f"Специальность: {SPECIALTY}",
        f"Модуль: {MODULE}",
        "",
        f"Дата: {TODAY}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).font.size = Pt(12)

    doc.add_page_break()

    # 1. Ссылка на репозиторий
    doc.add_heading("1. Ссылка на репозиторий", level=1)
    doc.add_paragraph(f"Репозиторий: {REPO}")
    doc.add_paragraph("Ветка: main")
    doc.add_paragraph(
        "В репозиторий добавлены файлы этапа 2: Dockerfile, docker-compose.yml, "
        ".dockerignore, Makefile, .editorconfig, INSTALL.md, BAT-скрипты в scripts/, "
        "настройки ESLint и Prettier."
    )

    # 2. Локальный запуск
    doc.add_heading("2. Как запустить проект локально", level=1)
    doc.add_paragraph("Windows:")
    for c in ["scripts\\setup.bat", "scripts\\run.bat"]:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph("macOS / Linux:")
    for c in ["make setup", "make run"]:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph("Адрес приложения: http://localhost:4000")
    doc.add_paragraph("Вход: admin@calipso.coffee / admin123")

    add_table(
        doc,
        ["Действие", "BAT", "Makefile", "npm"],
        [
            ["Установка", "scripts\\setup.bat", "make setup", "npm install"],
            ["Запуск", "scripts\\run.bat", "make run", "npm run dev"],
            ["Проверка", "scripts\\check.bat", "make check", "npm run check"],
            ["Формат", "scripts\\format.bat", "make format", "npm run format"],
        ],
    )

    # 3. Docker
    doc.add_heading("3. Как запустить через Docker", level=1)
    doc.add_paragraph("Перед запуском: cp .env.example .env")
    for c in [
        "docker compose up --build",
        "make docker-up",
        "scripts\\docker-up.bat",
    ]:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph("Адрес: http://localhost:4000")
    doc.add_paragraph("Остановка: docker compose down / make docker-down / scripts\\docker-down.bat")
    doc.add_paragraph("Логи: docker compose logs -f / make logs / scripts\\logs.bat")
    doc.add_paragraph(
        "Архитектура Docker: один сервис app (Node.js + Express + Prisma + SQLite). "
        "База данных хранится в volume calipso_data (путь file:/data/dev.db). "
        "docker-compose.yml используется, так как проект содержит backend и БД."
    )

    # 4. Инструменты качества
    doc.add_heading("4. Какие инструменты качества настроены", level=1)
    add_table(
        doc,
        ["Инструмент", "Назначение", "Команда"],
        [
            ["ESLint 8", "Линтер JavaScript (src/, public/, test/)", "npm run lint"],
            ["Prettier 3", "Форматирование кода", "npm run format"],
            ["node:test", "Unit-тесты (8 тестов)", "npm test"],
            ["npm run check", "Линтер + тесты одной командой", "make check / scripts\\check.bat"],
            [".editorconfig", "Единый стиль в редакторе", "—"],
        ],
    )

    # 5. Проблемы
    doc.add_heading("5. Какие проблемы были исправлены", level=1)
    add_table(
        doc,
        ["№", "Проблема", "Решение"],
        [
            [
                "1",
                "Проект запускался только вручную через IDE",
                "Добавлены Makefile, BAT-скрипты, INSTALL.md, команды setup/run/check/format",
            ],
            [
                "2",
                "Отсутствовали Docker и изолированный запуск",
                "Созданы Dockerfile, docker-compose.yml, .dockerignore, entrypoint-скрипт",
            ],
            [
                "3",
                "Не было линтера и форматтера",
                "Настроены ESLint + Prettier, команды npm run lint/format/check",
            ],
            [
                "4",
                "ESLint warning: неиспользуемый import в reports.routes.js",
                "Удалён лишний require('express')",
            ],
            [
                "5",
                "EADDRINUSE порт 4000 при повторном запуске",
                "Документировано в INSTALL.md: kill $(lsof -t -i :4000)",
            ],
        ],
    )

    # 6. Скриншоты
    doc.add_heading("6. Скриншоты", level=1)
    add_table(
        doc,
        ["Файл", "Что показано"],
        [
            ["01_clean_clone.png", "Проект клонирован в новую папку"],
            ["02_setup_bat_success.png", "Успешный scripts\\setup.bat или make setup"],
            ["03_make_help_or_make_check.png", "make help или make check без ошибок"],
            ["04_docker_build_success.png", "docker compose build — успешная сборка"],
            ["05_docker_compose_up.png", "docker compose up — сервис app запущен"],
            ["06_app_opened_in_browser.png", "http://localhost:4000 после входа"],
            ["07_linter_success.png", "npm run lint или make check"],
            ["08_formatter_success.png", "npm run format или make format"],
            ["09_git_commit.png", "git commit + push с файлами этапа 2"],
        ],
    )

    # Файлы в репозитории
    doc.add_heading("7. Добавленные файлы в репозитории", level=1)
    for f in [
        "Dockerfile",
        "docker-compose.yml",
        ".dockerignore",
        "Makefile",
        ".editorconfig",
        ".eslintrc.cjs",
        ".eslintignore",
        ".prettierrc",
        ".prettierignore",
        "INSTALL.md",
        "scripts/setup.bat",
        "scripts/run.bat",
        "scripts/check.bat",
        "scripts/format.bat",
        "scripts/docker-up.bat",
        "scripts/docker-down.bat",
        "scripts/logs.bat",
        "scripts/docker-entrypoint.sh",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph(f"Студент: {STUDENT_FIO}, группа {GROUP}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Дата: {TODAY}").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    OUT.mkdir(parents=True, exist_ok=True)
    (OUT / "docs").mkdir(exist_ok=True)
    doc.save(OUT / "docs" / "01_Короткий_отчет_по_этапу_2.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    (OUT / "screenshots").mkdir(exist_ok=True)
    (OUT / "docs").mkdir(exist_ok=True)

    (OUT / "repo_link.txt").write_text(REPO + "\n", encoding="utf-8")

    (OUT / "screenshots" / "README_скриншоты.txt").write_text(
        """Обязательные скриншоты этапа 2:

01_clean_clone.png — git clone в новую папку
02_setup_bat_success.png — scripts\\setup.bat или make setup
03_make_help_or_make_check.png — make help или make check
04_docker_build_success.png — docker compose build (Docker Desktop должен быть запущен)
05_docker_compose_up.png — docker compose up --build
06_app_opened_in_browser.png — браузер http://localhost:4000
07_linter_success.png — npm run lint / make check
08_formatter_success.png — npm run format / make format
09_git_commit.png — git commit и push
""",
        encoding="utf-8",
    )

    (OUT / "README.md").write_text(
        f"""# УП.03 Этап 2 — Calipso.Coffee

**Студент:** {STUDENT_FIO}  
**Группа:** {GROUP}

## Состав папки сдачи

| Элемент | Путь |
|---------|------|
| Ссылка на репозиторий | repo_link.txt |
| Короткий отчёт | docs/01_Короткий_отчет_по_этапу_2.docx |
| Скриншоты | screenshots/ (9 файлов) |
| Файлы проекта | в репозитории GitHub |

## Репозиторий

{REPO}

## Команды этапа 2

```bash
make setup && make run          # локально
make check && make format       # качество кода
make docker-up                  # Docker
```
""",
        encoding="utf-8",
    )

    build_report()
    print(f"Готово: {OUT}")


if __name__ == "__main__":
    main()
