#!/usr/bin/env python3
"""Генерация материалов УП.03 Этап 6 для Calipso.Coffee."""

import shutil
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

BASE = Path(__file__).resolve().parent.parent
OUT = BASE / "УП03_Этап6_Васильев_Данил_РПО-23-1"
REPO = "https://github.com/Aizenrix/Practicable.git"
PUBLIC_URL = "http://localhost:4000"
BRANCH = "support/fix-invalid-token-session"
TODAY = date.today().strftime("%d.%m.%Y")

STUDENT_FIO = "Васильев Данил Вячеславович"
GROUP = "РПО 23/1"


def build_support_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SUPPORT REPORT — Этап 6\nУП.03. CI/CD, релиз и инцидент поддержки\n")
    r.bold = True
    r.font.size = Pt(14)

    for text in [
        "Проект: Calipso.Coffee",
        f"Студент: {STUDENT_FIO}, группа {GROUP}",
        f"Дата: {TODAY}",
        f"Репозиторий: {REPO}",
        f"Ветка: {BRANCH}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)

    doc.add_page_break()

    doc.add_heading("1. Ссылка на репозиторий", level=1)
    doc.add_paragraph(f"{REPO}, ветка {BRANCH}, Pull Request в main.")

    doc.add_heading("2. Выбранная проблема (BUG-06)", level=1)
    doc.add_paragraph(
        "При восстановлении сессии с невалидным JWT в localStorage в DevTools Console "
        "появлялась красная ошибка «Неверный токен», хотя API корректно возвращает 401. "
        "Инцидент взят из DEFECT_LOG этапа 4, связан с риском эксплуатации этапа 5."
    )

    doc.add_heading("3. Воспроизведение", level=1)
    for step in [
        "Войти на http://localhost:4000 как admin@calipso.coffee",
        "Изменить token в localStorage на невалидное значение",
        "Обновить страницу (F5)",
        "До fix: Console — Error: Неверный токен",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("4. Диагностика", level=1)
    doc.add_paragraph(
        "Network: GET /api/auth/me → 401. Docker logs: GET /api/auth/me 401. "
        "Причина: public/app.js вызывал console.error для ожидаемого 401. "
        "Подробно: INCIDENT_REPORT.md, reports/incident_diagnostics.txt."
    )

    doc.add_heading("5. Изменения", level=1)
    for item in [
        "public/app.js — clearSession(), isExpectedAuthError()",
        ".github/workflows/ci.yml — GitHub Actions",
        "CHANGELOG.md, RELEASE_NOTES.md v1.0.1",
        "Makefile: build, release-check",
        "scripts: build.bat, release-check.bat, create-release.bat",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. Проверка", level=1)
    for item in [
        "make check",
        "make release-check",
        "GitHub Actions CI (push/PR)",
        "Ручная проверка: localStorage token → F5 без console.error",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Релиз", level=1)
    doc.add_paragraph("Версия 1.0.1, CHANGELOG.md, RELEASE_NOTES.md, тег v1.0.1, release/project_release.zip.")

    doc.add_heading("8. Скриншоты", level=1)
    for item in [
        "01_github_issue_created.png",
        "02_branch_created.png",
        "03_bug_reproduced.png",
        "04_logs_diagnostics.png",
        "05_fix_commit.png",
        "06_tests_passed_locally.png",
        "07_github_actions_success.png",
        "08_pull_request.png",
        "09_changelog_release_notes.png",
        "10_release_tag_or_archive.png",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph(f"Студент: {STUDENT_FIO}, группа {GROUP}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    (OUT / "docs").mkdir(parents=True, exist_ok=True)
    doc.save(OUT / "docs" / "SUPPORT_REPORT.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    for sub in ["docs", "reports", "screenshots", "project_files_in_repository"]:
        (OUT / sub).mkdir(parents=True, exist_ok=True)

    (OUT / "repo_link.txt").write_text(REPO + "\n", encoding="utf-8")
    (OUT / "branch.txt").write_text(BRANCH + "\n", encoding="utf-8")

    for name in ["INCIDENT_REPORT.md", "CHANGELOG.md", "RELEASE_NOTES.md", "RELEASE_CHECKLIST.md"]:
        src = BASE / name
        if src.exists():
            shutil.copy(src, OUT / "docs" / name)

    diag = BASE / "reports" / "incident_diagnostics.txt"
    if diag.exists():
        shutil.copy(diag, OUT / "reports" / "incident_diagnostics.txt")

    for f in [".github/workflows/ci.yml", "CHANGELOG.md", "RELEASE_NOTES.md"]:
        src = BASE / f
        if src.exists():
            dst = OUT / "project_files_in_repository" / Path(f).name
            if f.startswith(".github"):
                (OUT / "project_files_in_repository" / "ci.yml").write_bytes(src.read_bytes())
            else:
                shutil.copy(src, dst)

    for bat in ["test.bat", "build.bat", "release-check.bat", "create-release.bat"]:
        src = BASE / "scripts" / bat
        if src.exists():
            shutil.copy(src, OUT / "project_files_in_repository" / bat)

    (OUT / "screenshots" / "README_скриншоты.txt").write_text(
        """Скриншоты этапа 6 (10 шт.):

01_github_issue_created.png   — GitHub Issue BUG-06
02_branch_created.png         — git branch support/fix-invalid-token-session
03_bug_reproduced.png         — Console: Error Неверный токен (до fix)
04_logs_diagnostics.png       — Network 401 + docker logs
05_fix_commit.png             — git log / commit на ветке
06_tests_passed_locally.png   — make check / make release-check
07_github_actions_success.png — GitHub Actions CI green
08_pull_request.png           — Pull Request + Closes #issue
09_changelog_release_notes.png — CHANGELOG.md + RELEASE_NOTES.md
10_release_tag_or_archive.png — git tag v1.0.1 или release zip
""",
        encoding="utf-8",
    )

    (OUT / "README.md").write_text(
        f"""# УП.03 Этап 6 — Calipso.Coffee

**Студент:** {STUDENT_FIO}  
**Группа:** {GROUP}

## Инцидент

BUG-06: протухший JWT → красная ошибка в Console.

## Команды

```bash
make check
make build
make release-check
make build-release
```

## Состав

- docs/ — SUPPORT_REPORT.docx, INCIDENT_REPORT, CHANGELOG, RELEASE_NOTES, RELEASE_CHECKLIST
- screenshots/ — 10 скриншотов
- project_files_in_repository/ — ci.yml, bat-скрипты
""",
        encoding="utf-8",
    )

    build_support_report()
    print(f"Готово: {OUT}")


if __name__ == "__main__":
    main()
